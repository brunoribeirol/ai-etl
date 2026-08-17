"""E2E Scenario 4 — CSV + Postgres + REST + Document (PDF/DOCX) sources.

The 4th case-study scenario (Vault: `artefact/sprint-roadmap.md`, Sprint 5) —
adds `sources/document_source.py` (ADR-010) on top of Scenario 3's source mix,
proving the new connector wires through Orchestrator -> Extractor -> Transformer
-> Quality -> Loader the same way csv/postgres/rest already do.

Uses a real `.docx` file (via `python-docx`, so `document_source._extract_docx_text`
runs for real) and mocks only the LLM structuring call inside `document_source.py`
— same reasoning as `mock_pipeline_llm` for the other agent-level LLM calls.
"""

import json
from unittest.mock import MagicMock

import docx
import pandas as pd
import sqlalchemy

from ai_etl.services.execution_queue import enqueue_analysis, get_task_status
from tests.e2e.conftest import TEST_POSTGRES_URL, requires_full_stack

RETURNS_ROWS = [
    {"order_id": 1, "reason": "wrong size"},
    {"order_id": 3, "reason": "defective"},
    {"order_id": 4, "reason": "changed mind"},
    {"order_id": 5, "reason": "late delivery"},
    {"order_id": 6, "reason": "wrong item"},
]

TRANSFORM_CODE = """
def transform(dfs: dict) -> pd.DataFrame:
    orders = dfs["orders"]
    customers = dfs["customers"]
    returns = dfs["returns"]
    df = orders.merge(customers, on="customer_id", how="inner")
    df = df[df["status"] == "active"]
    df = df.merge(returns, on="order_id", how="left")
    return df
"""


@requires_full_stack
def test_scenario4_document_source_runs_end_to_end(
    tmp_path, mock_pipeline_llm, test_tenant, postgres_customers_table, mocker
) -> None:
    # 6 orders, 5 with a matching return reason (order 2 has none) — keeps the
    # left-joined "reason" column's null ratio at ~17%, under quality.py's 20%
    # NULL_ERROR_THRESHOLD, so this stays a "warning" (Loader still runs).
    orders_path = tmp_path / "orders.csv"
    pd.DataFrame(
        {
            "order_id": [1, 2, 3, 4, 5, 6],
            "customer_id": [100, 200, 300, 400, 500, 600],
            "status": ["active"] * 6,
        }
    ).to_csv(orders_path, index=False)

    # A real DOCX file, real python-docx extraction — the only mocked part is
    # the LLM call that structures the extracted text into rows.
    returns_doc_path = tmp_path / "returns.docx"
    document = docx.Document()
    document.add_paragraph("Order 1 was returned: wrong size.")
    document.add_paragraph("Order 3 was returned: defective.")
    document.add_paragraph("Order 4 was returned: changed mind.")
    document.add_paragraph("Order 5 was returned: late delivery.")
    document.add_paragraph("Order 6 was returned: wrong item.")
    document.save(str(returns_doc_path))

    document_llm = MagicMock()
    document_llm.invoke.return_value = MagicMock(content=json.dumps(RETURNS_ROWS))
    mocker.patch("ai_etl.sources.document_source.get_llm", return_value=document_llm)

    plan = {
        "sources": [
            {"name": "orders", "type": "csv", "path": str(orders_path)},
            {"name": "customers", "type": "postgres", "table": "public.customers"},
            {"name": "returns", "type": "document", "path": str(returns_doc_path)},
        ],
        "destination": {"type": "postgres", "table": "public.enriched_orders"},
        "transformations": [
            "join orders and customers on customer_id",
            'filter status == "active"',
            "left join returns on order_id",
        ],
        "quality_checks": ["null_check"],
    }
    mock_pipeline_llm(plan, TRANSFORM_CODE)

    spec = (
        "Read orders.csv, public.customers, and returns.docx. Join orders and "
        "customers on customer_id, filter active orders, left join the "
        "returns document on order_id. Save to public.enriched_orders."
    )
    tenant_id = test_tenant["tenant_id"]

    task_id = enqueue_analysis(
        spec, business_question="", run_dir=str(tmp_path), tenant_id=tenant_id
    )
    status = get_task_status(task_id)

    assert status["state"] == "SUCCESS", status.get("error")
    assert status["result"]["status"] == "completed"

    engine = sqlalchemy.create_engine(TEST_POSTGRES_URL)
    written = pd.read_sql("SELECT * FROM public.enriched_orders", engine)
    assert len(written) == 6
    order_1 = written[written["order_id"] == 1].iloc[0]
    assert order_1["reason"] == "wrong size"
    order_2 = written[written["order_id"] == 2].iloc[0]
    assert pd.isna(order_2["reason"])
