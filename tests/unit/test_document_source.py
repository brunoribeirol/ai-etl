"""Unit tests for the document (PDF/DOCX) source connector (ADR-010)."""

import json
from unittest.mock import MagicMock

import pandas as pd
import pytest

from ai_etl.sources.document_source import (
    MAX_ATTEMPTS,
    _extract_docx_text,
    _structure_text,
    extract_document_text,
    load_document,
)

VALID_ROWS = [
    {"product": "A", "revenue": 100},
    {"product": "B", "revenue": 200},
]


def _mock_llm(responses: list[str]) -> MagicMock:
    llm = MagicMock()
    llm.invoke.side_effect = [MagicMock(content=r) for r in responses]
    return llm


@pytest.fixture
def mock_get_llm(mocker):
    return mocker.patch("ai_etl.sources.document_source.get_llm")


def test_extract_document_text_rejects_unsupported_extension() -> None:
    with pytest.raises(ValueError, match="Unsupported document type"):
        extract_document_text("report.txt")


def test_extract_document_text_dispatches_pdf(mocker) -> None:
    mock_pdf = mocker.patch("ai_etl.sources.document_source._extract_pdf_text", return_value="x")
    extract_document_text("report.pdf")
    mock_pdf.assert_called_once_with("report.pdf")


def test_extract_document_text_dispatches_docx(mocker) -> None:
    mock_docx = mocker.patch("ai_etl.sources.document_source._extract_docx_text", return_value="x")
    extract_document_text("report.docx")
    mock_docx.assert_called_once_with("report.docx")


def test_extract_docx_text_includes_table_cells(tmp_path) -> None:
    """Regression test for the Data Engineer 2026-08-24 audit finding: a real
    .docx with a heading + a real table previously vanished entirely because
    `_extract_docx_text` only read `document.paragraphs`. Builds a real .docx
    on disk (no mocking of python-docx) so the table-reading code path is
    actually exercised."""
    import docx

    document = docx.Document()
    document.add_heading("Quarterly Revenue", level=1)

    table = document.add_table(rows=4, cols=4)
    header = ["Region", "Q1", "Q2", "Q3"]
    rows = [
        ["North", "100", "110", "120"],
        ["South", "200", "210", "220"],
        ["East", "300", "310", "320"],
    ]
    for col_idx, header_text in enumerate(header):
        table.cell(0, col_idx).text = header_text
    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, cell_value in enumerate(row_values):
            table.cell(row_idx, col_idx).text = cell_value

    docx_path = tmp_path / "report.docx"
    document.save(str(docx_path))

    text = _extract_docx_text(str(docx_path))

    assert "Quarterly Revenue" in text
    assert "Region | Q1 | Q2 | Q3" in text
    assert "North | 100 | 110 | 120" in text
    assert "South | 200 | 210 | 220" in text
    assert "East | 300 | 310 | 320" in text
    assert "[Table]" in text
    assert "[/Table]" in text


def test_structure_text_valid_json_returns_dataframe(mock_get_llm) -> None:
    mock_get_llm.return_value = _mock_llm([json.dumps(VALID_ROWS)])
    df = _structure_text("some extracted document text")

    assert isinstance(df, pd.DataFrame)
    assert list(df.columns) == ["product", "revenue"]
    assert len(df) == 2


def test_structure_text_retries_on_invalid_json(mock_get_llm) -> None:
    mock_get_llm.return_value = _mock_llm(["not json at all", json.dumps(VALID_ROWS)])
    df = _structure_text("some text")

    assert len(df) == 2


def test_structure_text_retries_on_non_array_json(mock_get_llm) -> None:
    """A JSON object (not an array) is syntactically valid JSON but the wrong
    shape — must be treated as a retry-able failure, not accepted as-is."""
    mock_get_llm.return_value = _mock_llm([json.dumps({"not": "a list"}), json.dumps(VALID_ROWS)])
    df = _structure_text("some text")

    assert len(df) == 2


def test_structure_text_exhausts_retries_and_raises(mock_get_llm) -> None:
    mock_get_llm.return_value = _mock_llm(["not json"] * MAX_ATTEMPTS)

    with pytest.raises(ValueError, match="Failed to extract structured data"):
        _structure_text("some text")

    assert mock_get_llm.return_value.invoke.call_count == MAX_ATTEMPTS


def test_load_document_composes_extract_and_structure(mocker) -> None:
    mocker.patch("ai_etl.sources.document_source.extract_document_text", return_value="raw text")
    mock_structure = mocker.patch(
        "ai_etl.sources.document_source._structure_text",
        return_value=pd.DataFrame(VALID_ROWS),
    )

    df = load_document("report.pdf")

    mock_structure.assert_called_once_with("raw text", None, None)
    assert len(df) == 2


def test_load_document_forwards_llm_override(mocker) -> None:
    """Sprint 30/gap-closing (ADR-031 §5) — a per-pipeline LLM override reaches
    document_source.py's own get_llm() call the same way it reaches every other
    call site."""
    mocker.patch("ai_etl.sources.document_source.extract_document_text", return_value="raw text")
    mock_structure = mocker.patch(
        "ai_etl.sources.document_source._structure_text",
        return_value=pd.DataFrame(VALID_ROWS),
    )

    load_document(
        "report.pdf", llm_provider_override="anthropic", llm_model_override="claude-haiku-4-5"
    )

    mock_structure.assert_called_once_with("raw text", "anthropic", "claude-haiku-4-5")


def test_structure_text_forwards_override_to_get_llm(mock_get_llm) -> None:
    mock_get_llm.return_value = _mock_llm([json.dumps(VALID_ROWS)])
    _structure_text("some text", "anthropic", "claude-haiku-4-5")

    mock_get_llm.assert_called_once_with(provider="anthropic", model="claude-haiku-4-5")
